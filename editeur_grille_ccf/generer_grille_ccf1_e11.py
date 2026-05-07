#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Generateur Grille CCF1 E1.1 (CAPa Jardinier Paysagiste)
========================================================
Remplit le modele Word officiel "Grille_CCF_vierge.docx" a partir du JSON
exporte par ccf1_e11.html (bouton "JSON eleve").

Le modele n'est PAS reconstruit : on ne fait qu'ecrire dans les cellules
existantes. La mise en page reste 100% identique au modele.

Usage:
    python generer_grille_ccf1_e11.py eleve.json
    python generer_grille_ccf1_e11.py eleve.json --template autre_modele.docx
    python generer_grille_ccf1_e11.py eleve.json --output Grille_DUPONT.docx
"""

import argparse
import json
import os
import sys
from copy import deepcopy

from docx import Document

HERE = os.path.dirname(os.path.abspath(__file__))
DEFAULT_TEMPLATE = os.path.join(HERE, "Grille_CCF_vierge.docx")

# Mapping (table_index, row, col) du modele actuel.
# Verifie sur "Grille CCF vierge.docx" (3 tables, voir README).
T_HEADER = 2  # Table 2 = grille principale

# Lignes par phase (premier indicateur, nb d'indicateurs)
PHASES = {
    1: {"start_row": 4,  "n": 3, "bareme": 5},
    2: {"start_row": 7,  "n": 4, "bareme": 6},
    3: {"start_row": 11, "n": 3, "bareme": 6},
    4: {"start_row": 14, "n": 2, "bareme": 3},
}

# Colonnes
COL_CRITERE   = 0
COL_INDIC     = 1
COL_EVAL_BASE = 2  # --, -, +, ++  -> 2,3,4,5
COL_NOTE      = 6
COL_OBS       = 7

NIVEAU_TO_COL = {"--": 2, "-": 3, "+": 4, "++": 5}
VAL_TO_COL    = {0: 2, 1: 3, 2: 4, 3: 5}


def set_cell_text(cell, text, keep_format=True):
    """Remplace le texte d'une cellule en conservant la mise en forme du
    premier run du premier paragraphe. Les autres paragraphes/runs sont vides."""
    paragraphs = cell.paragraphs
    if not paragraphs:
        cell.add_paragraph(text)
        return
    p0 = paragraphs[0]
    if p0.runs:
        # garder le premier run (style), vider les suivants
        p0.runs[0].text = text
        for r in p0.runs[1:]:
            r.text = ""
    else:
        p0.add_run(text)
    # vider les paragraphes suivants
    for p in paragraphs[1:]:
        for r in p.runs:
            r.text = ""


def append_cell_line(cell, text):
    """Ajoute un paragraphe sans casser le contenu existant."""
    cell.add_paragraph(text)


def cocher(cell):
    """Place un X centre dans la cellule en gardant la mise en forme."""
    set_cell_text(cell, "X")
    # tentative de centrage
    try:
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        if cell.paragraphs[0].runs:
            cell.paragraphs[0].runs[0].bold = True
    except Exception:
        pass


def ecrire_note(cell, note, bareme):
    """Ecrit la note sous forme 'note/bareme' en preservant le style."""
    if note is None or note == "":
        txt = "/{}".format(bareme)
    else:
        # accepte int ou float
        if isinstance(note, float) and note.is_integer():
            note = int(note)
        txt = "{}/{}".format(note, bareme)
    set_cell_text(cell, txt)


def remplir(doc, data):
    """data = dict racine du JSON exporte (avec cle 'eleve')."""
    eleve = data.get("eleve", data)  # tolere un JSON deja a plat
    nom        = eleve.get("nom", "")
    date       = eleve.get("date", "")
    sujet      = eleve.get("sujet", "") or eleve.get("sujet_n", "")
    apprec     = eleve.get("appreciation_generale", "")
    total      = eleve.get("total")
    sig_esc    = eleve.get("enseignant_esc", "")
    sig_hg     = eleve.get("enseignant_hg", "")
    phases     = eleve.get("phases", {})

    grille = doc.tables[T_HEADER]

    # --- En-tete de la grille ---
    # R0[7] = Date
    if date:
        set_cell_text(grille.cell(0, 7), "Date : {}".format(date))
    # R1[0] = "Nom et Prenom : ..."
    set_cell_text(grille.cell(1, 0), "Nom et Prénom : {}".format(nom))
    # R1[2] = Sujet n°
    if sujet:
        set_cell_text(grille.cell(1, 2), "Sujet n° : {}".format(sujet))

    # --- Phases ---
    for p_num, conf in PHASES.items():
        phase = phases.get(str(p_num)) or phases.get(p_num) or {}
        indicateurs = phase.get("indicateurs", []) or []
        note        = phase.get("note")
        observations = phase.get("observations", "")
        bareme      = phase.get("bareme", conf["bareme"])

        # Cocher les niveaux
        for i, ind in enumerate(indicateurs):
            row = conf["start_row"] + i
            niveau = ind.get("niveau") if isinstance(ind, dict) else None
            valeur = ind.get("valeur") if isinstance(ind, dict) else None
            col = None
            if valeur is not None and valeur in VAL_TO_COL:
                col = VAL_TO_COL[valeur]
            elif niveau in NIVEAU_TO_COL:
                col = NIVEAU_TO_COL[niveau]
            if col is not None:
                cocher(grille.cell(row, col))

        # Note : seulement sur la premiere ligne de la phase
        ecrire_note(grille.cell(conf["start_row"], COL_NOTE), note, bareme)

        # Observations : sur la premiere ligne
        if observations:
            set_cell_text(grille.cell(conf["start_row"], COL_OBS), observations)

    # --- Derniere ligne : Appreciation et Total ---
    last = len(grille.rows) - 1  # 16
    if apprec:
        set_cell_text(grille.cell(last, 0), "APPRÉCIATION GÉNÉRALE : {}".format(apprec))
    if total is not None:
        if isinstance(total, float) and total.is_integer():
            total = int(total)
        set_cell_text(grille.cell(last, 7), "TOTAL : {} /20".format(total))

    # --- Signatures (ajoutees apres le tableau si presentes) ---
    if sig_esc or sig_hg:
        doc.add_paragraph("")
        p = doc.add_paragraph()
        p.add_run("Enseignant ESC : {}".format(sig_esc or "")).bold = True
        p2 = doc.add_paragraph()
        p2.add_run("Enseignant HG : {}".format(sig_hg or "")).bold = True


def main():
    ap = argparse.ArgumentParser(description="Remplit la grille CCF1 E1.1 a partir d'un JSON eleve.")
    ap.add_argument("json_file", help="Fichier JSON exporte par ccf1_e11.html")
    ap.add_argument("--template", default=DEFAULT_TEMPLATE, help="Modele .docx a remplir")
    ap.add_argument("--output", default=None, help="Fichier de sortie")
    args = ap.parse_args()

    if not os.path.exists(args.template):
        sys.exit("Modele introuvable : {}".format(args.template))
    if not os.path.exists(args.json_file):
        sys.exit("JSON introuvable : {}".format(args.json_file))

    with open(args.json_file, "r", encoding="utf-8") as f:
        data = json.load(f)

    doc = Document(args.template)
    remplir(doc, data)

    if args.output:
        out = args.output
    else:
        nom = (data.get("eleve", {}).get("nom") or "eleve").strip().replace(" ", "_")
        out = os.path.join(os.getcwd(), "Grille_CCF1_E11_{}.docx".format(nom))

    doc.save(out)
    print("OK ->", out)


if __name__ == "__main__":
    main()
