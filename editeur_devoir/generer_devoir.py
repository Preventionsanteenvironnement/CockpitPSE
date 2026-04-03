#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Generateur de Devoirs PSE - Document Word standardise
=====================================================
Convertit un fichier JSON (genere par le prompt IA) en document Word (.docx)
avec une presentation standardisee : en-tete, 3 approches, grille de notation,
et corrige optionnel.

Usage:
    python generer_devoir.py devoir.json
    python generer_devoir.py devoir.json --output mon_devoir.docx
"""

import json
import sys
import os
from datetime import datetime

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# =============================================================================
# COULEURS
# =============================================================================
BLEU_FONCE = RGBColor(0x1A, 0x52, 0x76)
BLEU_MOYEN = RGBColor(0x29, 0x80, 0xB9)
BLEU_CLAIR = RGBColor(0x34, 0x98, 0xDB)
BLEU_PALE = RGBColor(0xEB, 0xF5, 0xFB)
VERT = RGBColor(0x27, 0xAE, 0x60)
VERT_PALE = RGBColor(0xD5, 0xF5, 0xE3)
GRIS = RGBColor(0x55, 0x55, 0x55)
GRIS_CLAIR = RGBColor(0xDD, 0xDD, 0xDD)
GRIS_TRES_CLAIR = RGBColor(0xEE, 0xEE, 0xEE)
ROUGE = RGBColor(0xE7, 0x4C, 0x3C)
NOIR = RGBColor(0x00, 0x00, 0x00)
BLANC = RGBColor(0xFF, 0xFF, 0xFF)

# Couleurs par competence
COULEURS_COMPETENCES = {
    "C1": RGBColor(0xE7, 0x4C, 0x3C),  # Rouge
    "C2": RGBColor(0x34, 0x98, 0xDB),  # Bleu
    "C3": RGBColor(0x2E, 0xCC, 0x71),  # Vert
    "C4": RGBColor(0xF3, 0x9C, 0x12),  # Orange
    "C5": RGBColor(0x9B, 0x59, 0xB6),  # Violet
    "C6": RGBColor(0x1A, 0xBC, 0x9C),  # Turquoise
}

FOND_HEX = {
    "bleu_pale": "EBF5FB",
    "vert_pale": "D5F5E3",
    "gris_clair": "EEEEEE",
    "bleu_fonce": "1A5276",
    "blanc": "FFFFFF",
    "rouge_pale": "FDEDEC",
    "orange_pale": "FEF9E7",
}


# =============================================================================
# UTILITAIRES
# =============================================================================

def set_cell_shading(cell, color_hex):
    """Applique une couleur de fond a une cellule."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell_border(cell, **kwargs):
    """Definit les bordures d'une cellule.
    kwargs: top, bottom, left, right avec valeurs dict(sz, color, val)
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}></w:tcBorders>')
    for edge, attrs in kwargs.items():
        element = parse_xml(
            f'<w:{edge} {nsdecls("w")} w:val="{attrs.get("val", "single")}" '
            f'w:sz="{attrs.get("sz", "4")}" w:space="0" '
            f'w:color="{attrs.get("color", "000000")}"/>'
        )
        tcBorders.append(element)
    tcPr.append(tcBorders)


def add_formatted_run(paragraph, text, bold=False, italic=False, size=12,
                      color=None, font_name="Arial"):
    """Ajoute un run formate a un paragraphe."""
    run = paragraph.add_run(text)
    run.font.name = font_name
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color
    return run


def add_paragraph_formatted(doc, text, bold=False, italic=False, size=12,
                            color=None, alignment=None, space_before=0,
                            space_after=6, font_name="Arial"):
    """Ajoute un paragraphe formate au document."""
    p = doc.add_paragraph()
    if alignment is not None:
        p.alignment = alignment
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    add_formatted_run(p, text, bold=bold, italic=italic, size=size,
                      color=color, font_name=font_name)
    return p


def set_table_width(table, width_inches):
    """Definit la largeur d'un tableau."""
    for row in table.rows:
        for cell in row.cells:
            cell.width = Inches(width_inches / len(row.cells))


# =============================================================================
# CONSTRUCTION DU DOCUMENT
# =============================================================================

def build_header(doc, meta):
    """Construit l'en-tete du devoir."""
    # Titre principal
    add_paragraph_formatted(
        doc, "Prevention Sante Environnement",
        bold=True, size=18, color=BLEU_FONCE,
        alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=4
    )

    # Sous-titre
    add_paragraph_formatted(
        doc, meta["titre"],
        bold=True, size=14, color=BLEU_MOYEN,
        alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=12
    )

    # Tableau infos eleve
    table = doc.add_table(rows=2, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # En-tete du tableau
    headers = ["Nom", "Prenom", "Classe", "Date"]
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, FOND_HEX["bleu_fonce"])
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_formatted_run(p, header, bold=True, size=11, color=BLANC)

    # Lignes vides pour remplissage
    for i in range(4):
        cell = table.rows[1].cells[i]
        set_cell_shading(cell, FOND_HEX["bleu_pale"])
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(8)
        add_formatted_run(p, " ", size=12)

    doc.add_paragraph()  # Espace

    # Cadre informations devoir
    info_table = doc.add_table(rows=1, cols=4)
    info_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    infos = [
        ("Niveau", meta.get("niveau", "")),
        ("Duree", meta.get("duree", "")),
        ("Difficulte", meta.get("difficulte", "").capitalize()),
        ("Bareme", f'/ {meta.get("bareme_total", 20)}'),
    ]

    for i, (label, value) in enumerate(infos):
        cell = info_table.rows[0].cells[i]
        set_cell_shading(cell, FOND_HEX["bleu_pale"])
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_formatted_run(p, f"{label} : ", bold=True, size=10, color=BLEU_FONCE)
        add_formatted_run(p, value, size=10, color=NOIR)

    doc.add_paragraph()

    # Consigne generale
    if meta.get("consigne_generale"):
        add_paragraph_formatted(
            doc, meta["consigne_generale"],
            italic=True, size=11, color=GRIS,
            alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=12
        )

    # Ligne separatrice
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run("_" * 80)
    run.font.color.rgb = GRIS_CLAIR
    run.font.size = Pt(8)


def build_situation(doc, situation):
    """Construit l'encadre de situation."""
    # Titre situation
    add_paragraph_formatted(
        doc, situation.get("titre", "Situation"),
        bold=True, size=12, color=BLEU_FONCE, space_before=6, space_after=6
    )

    # Encadre de situation (tableau 1 cellule)
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.rows[0].cells[0]
    set_cell_shading(cell, FOND_HEX["bleu_pale"])
    set_cell_border(
        cell,
        top={"sz": "12", "color": "2980B9", "val": "single"},
        bottom={"sz": "12", "color": "2980B9", "val": "single"},
        left={"sz": "12", "color": "2980B9", "val": "single"},
        right={"sz": "12", "color": "2980B9", "val": "single"},
    )

    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)
    add_formatted_run(p, situation.get("texte", ""), size=11)

    doc.add_paragraph()  # Espace apres situation


def build_document_ref(doc, document):
    """Construit un encadre document de reference."""
    # En-tete document
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.rows[0].cells[0]
    set_cell_shading(cell, FOND_HEX["vert_pale"])

    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    add_formatted_run(p, document.get("titre", "Document"), bold=True, size=11,
                      color=BLEU_FONCE)

    # Contenu du document
    if document.get("contenu"):
        table2 = doc.add_table(rows=1, cols=1)
        table2.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell2 = table2.rows[0].cells[0]
        set_cell_border(
            cell2,
            top={"sz": "4", "color": "CCCCCC"},
            bottom={"sz": "4", "color": "CCCCCC"},
            left={"sz": "4", "color": "CCCCCC"},
            right={"sz": "4", "color": "CCCCCC"},
        )
        p2 = cell2.paragraphs[0]
        p2.paragraph_format.space_before = Pt(6)
        p2.paragraph_format.space_after = Pt(6)
        add_formatted_run(p2, document["contenu"], size=10)

    # Espace
    p_space = doc.add_paragraph()
    p_space.paragraph_format.space_before = Pt(2)
    p_space.paragraph_format.space_after = Pt(4)


def build_question(doc, question, is_corrige=False):
    """Construit une question avec son espace de reponse."""
    comp = question.get("competence", "")
    comp_color = COULEURS_COMPETENCES.get(comp, NOIR)

    # Ligne de question : numero + competence + texte
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)

    # Numero
    add_formatted_run(p, f'{question["numero"]}) ', bold=True, size=12,
                      color=BLEU_FONCE)
    # Competence
    add_formatted_run(p, f'[{comp}] ', bold=True, size=11, color=comp_color)
    # Texte question
    add_formatted_run(p, question.get("texte", ""), size=11)

    q_type = question.get("type", "reponse_courte")

    if is_corrige:
        _build_corrige_answer(doc, question)
    else:
        _build_answer_space(doc, question, q_type)


def _build_answer_space(doc, question, q_type):
    """Construit l'espace de reponse selon le type de question."""

    if q_type in ("qcm", "cocher"):
        _build_qcm(doc, question, multi=False)

    elif q_type == "qcm_multi":
        _build_qcm(doc, question, multi=True)

    elif q_type in ("vrai_faux", "vrai_faux_justifie"):
        _build_vrai_faux(doc, question)

    elif q_type == "tableau":
        _build_tableau(doc, question)

    elif q_type == "texte_trous":
        _build_texte_trous(doc, question)

    elif q_type == "classement":
        _build_classement(doc, question)

    elif q_type == "appariement":
        _build_appariement(doc, question)

    elif q_type == "mise_en_ordre":
        _build_mise_en_ordre(doc, question)

    elif q_type == "grille_risques":
        _build_grille_risques(doc, question)

    elif q_type == "schema":
        _build_schema(doc, question)

    elif q_type in ("relever", "reponse_courte", "calcul"):
        _build_lignes_reponse(doc, question, nb_lignes=3)

    elif q_type == "redaction":
        _build_lignes_reponse(doc, question, nb_lignes=8)

    else:
        _build_lignes_reponse(doc, question, nb_lignes=3)


def _build_qcm(doc, question, multi=False):
    """Construit un QCM."""
    instruction = "(Cochez la ou les bonne(s) reponse(s))" if multi else "(Cochez la bonne reponse)"
    add_paragraph_formatted(doc, instruction, italic=True, size=9, color=GRIS,
                            space_before=2, space_after=4)

    options = question.get("options", [])
    for opt in options:
        label = opt if isinstance(opt, str) else opt.get("label", "")
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.left_indent = Cm(1)
        # Case a cocher Unicode
        add_formatted_run(p, "   \u2610   ", size=14)
        add_formatted_run(p, label, size=11)


def _build_vrai_faux(doc, question):
    """Construit un exercice Vrai/Faux."""
    affirmations = question.get("affirmations", [])
    if not affirmations:
        # Simple V/F
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(1)
        add_formatted_run(p, "   \u2610 Vrai        \u2610 Faux", size=12)
        if question.get("type") == "vrai_faux_justifie":
            add_paragraph_formatted(doc, "Justification :", bold=True, size=10,
                                    space_before=4)
            _add_dotted_lines(doc, 3)
    else:
        # Tableau V/F avec plusieurs affirmations
        nb_rows = len(affirmations) + 1
        cols = 3 if question.get("type") != "vrai_faux_justifie" else 4
        table = doc.add_table(rows=nb_rows, cols=cols)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # En-tetes
        headers = ["Affirmation", "Vrai", "Faux"]
        if cols == 4:
            headers.append("Justification")

        for i, h in enumerate(headers):
            cell = table.rows[0].cells[i]
            set_cell_shading(cell, FOND_HEX["bleu_fonce"])
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_formatted_run(p, h, bold=True, size=10, color=BLANC)

        for idx, aff in enumerate(affirmations):
            row = table.rows[idx + 1]
            text = aff if isinstance(aff, str) else aff.get("texte", "")

            # Affirmation
            p = row.cells[0].paragraphs[0]
            add_formatted_run(p, text, size=10)

            # Cases V/F
            for col_idx in (1, 2):
                p = row.cells[col_idx].paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                add_formatted_run(p, "\u2610", size=14)

            # Justification
            if cols == 4:
                p = row.cells[3].paragraphs[0]
                add_formatted_run(p, ".....................", size=9, color=GRIS_CLAIR)

            # Alternance couleur
            if idx % 2 == 0:
                for c in row.cells:
                    set_cell_shading(c, FOND_HEX["bleu_pale"])

    doc.add_paragraph()


def _build_tableau(doc, question):
    """Construit un tableau a completer."""
    tableau = question.get("tableau", {})
    colonnes = tableau.get("colonnes", ["Element 1", "Element 2"])
    nb_lignes = tableau.get("lignes_vides", 2)

    table = doc.add_table(rows=1 + nb_lignes, cols=len(colonnes))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # En-tetes
    for i, col_name in enumerate(colonnes):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, FOND_HEX["bleu_fonce"])
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_formatted_run(p, col_name, bold=True, size=10, color=BLANC)

    # Lignes vides
    for r in range(1, nb_lignes + 1):
        for c in range(len(colonnes)):
            cell = table.rows[r].cells[c]
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(12)
            if r % 2 == 0:
                set_cell_shading(cell, FOND_HEX["bleu_pale"])

    # Aide
    aide = tableau.get("aide")
    if aide:
        add_paragraph_formatted(doc, f"Aide : {aide}", italic=True, size=9,
                                color=GRIS, space_before=4)
    doc.add_paragraph()


def _build_texte_trous(doc, question):
    """Construit un texte a trous."""
    mots = question.get("mots_proposes", [])
    if mots:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        add_formatted_run(p, "Mots proposes : ", bold=True, size=10, color=BLEU_FONCE)
        add_formatted_run(p, " - ".join(mots), italic=True, size=10)

    texte = question.get("texte_a_trous", "")
    if texte:
        table = doc.add_table(rows=1, cols=1)
        cell = table.rows[0].cells[0]
        set_cell_shading(cell, FOND_HEX["bleu_pale"])
        set_cell_border(
            cell,
            top={"sz": "4", "color": "2980B9"},
            bottom={"sz": "4", "color": "2980B9"},
            left={"sz": "4", "color": "2980B9"},
            right={"sz": "4", "color": "2980B9"},
        )
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(8)
        p.paragraph_format.line_spacing = Pt(20)
        add_formatted_run(p, texte, size=11)
    doc.add_paragraph()


def _build_classement(doc, question):
    """Construit un exercice de classement."""
    categories = question.get("categories", [])
    elements = question.get("elements_a_classer", [])

    if elements:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        add_formatted_run(p, "Elements a classer : ", bold=True, size=10,
                          color=BLEU_FONCE)
        add_formatted_run(p, " | ".join(elements), size=10)

    if categories:
        table = doc.add_table(rows=2, cols=len(categories))
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        for i, cat in enumerate(categories):
            # En-tete categorie
            cell = table.rows[0].cells[i]
            set_cell_shading(cell, FOND_HEX["bleu_fonce"])
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_formatted_run(p, cat, bold=True, size=10, color=BLANC)

            # Case vide
            cell = table.rows[1].cells[i]
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(20)
            p.paragraph_format.space_after = Pt(20)

    doc.add_paragraph()


def _build_appariement(doc, question):
    """Construit un exercice d'appariement (relier)."""
    paires = question.get("paires", [])
    if not paires:
        _build_lignes_reponse(doc, question, nb_lignes=4)
        return

    table = doc.add_table(rows=len(paires) + 1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    headers = ["Terme", "", "Definition"]
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, FOND_HEX["bleu_fonce"])
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_formatted_run(p, h, bold=True, size=10, color=BLANC)

    for idx, paire in enumerate(paires):
        row = table.rows[idx + 1]
        # Terme
        p = row.cells[0].paragraphs[0]
        terme = paire.get("terme", "") if isinstance(paire, dict) else str(paire)
        add_formatted_run(p, terme, bold=True, size=10)

        # Fleche
        p = row.cells[1].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_formatted_run(p, "-->", size=10, color=GRIS)

        # Definition (vide pour l'eleve)
        p = row.cells[2].paragraphs[0]
        add_formatted_run(p, "..............................", size=9, color=GRIS_CLAIR)

        if idx % 2 == 0:
            for c in row.cells:
                set_cell_shading(c, FOND_HEX["bleu_pale"])

    # Liste des definitions melangees
    definitions = question.get("definitions_melangees", [])
    if definitions:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        add_formatted_run(p, "Definitions (dans le desordre) : ", bold=True,
                          size=10, color=BLEU_FONCE)
        for i, d in enumerate(definitions):
            p2 = doc.add_paragraph()
            p2.paragraph_format.left_indent = Cm(1)
            p2.paragraph_format.space_before = Pt(1)
            p2.paragraph_format.space_after = Pt(1)
            add_formatted_run(p2, f"{chr(65+i)}. ", bold=True, size=10)
            add_formatted_run(p2, d, size=10)

    doc.add_paragraph()


def _build_mise_en_ordre(doc, question):
    """Construit un exercice de mise en ordre."""
    etapes = question.get("etapes_melangees", [])
    if etapes:
        add_paragraph_formatted(
            doc, "(Numerotez les etapes dans l'ordre correct)",
            italic=True, size=9, color=GRIS, space_before=2
        )
        for etape in etapes:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1)
            p.paragraph_format.space_before = Pt(2)
            add_formatted_run(p, "  N\u00b0___  ", size=11, color=BLEU_MOYEN)
            add_formatted_run(p, etape, size=11)
    else:
        _build_lignes_reponse(doc, question, nb_lignes=4)
    doc.add_paragraph()


def _build_grille_risques(doc, question):
    """Construit une grille d'evaluation des risques."""
    colonnes = question.get("colonnes_grille",
                            ["Nature du risque", "Situation dangereuse",
                             "Evenement dangereux", "Dommage",
                             "Gravite (1-4)", "Probabilite (1-4)", "Priorite"])
    nb_lignes = question.get("lignes_vides", 2)

    table = doc.add_table(rows=1 + nb_lignes, cols=len(colonnes))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, col in enumerate(colonnes):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, FOND_HEX["bleu_fonce"])
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_formatted_run(p, col, bold=True, size=8, color=BLANC)

    for r in range(1, nb_lignes + 1):
        for c in range(len(colonnes)):
            cell = table.rows[r].cells[c]
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(10)

    doc.add_paragraph()


def _build_schema(doc, question):
    """Construit un schema/diagramme a completer (cases avec fleches)."""
    cases = question.get("cases", [])
    if not cases:
        _build_lignes_reponse(doc, question, nb_lignes=6)
        return

    # Schema sous forme de tableau horizontal
    nb_cases = len(cases)
    # Avec fleches entre les cases : case -> case -> case
    cols = nb_cases * 2 - 1
    table = doc.add_table(rows=2, cols=cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, case in enumerate(cases):
        col_idx = i * 2
        # Titre de la case
        cell = table.rows[0].cells[col_idx]
        set_cell_shading(cell, FOND_HEX["bleu_fonce"])
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        label = case if isinstance(case, str) else case.get("label", f"Case {i+1}")
        add_formatted_run(p, label, bold=True, size=9, color=BLANC)

        # Zone de reponse
        cell2 = table.rows[1].cells[col_idx]
        set_cell_border(
            cell2,
            top={"sz": "4", "color": "2980B9"},
            bottom={"sz": "4", "color": "2980B9"},
            left={"sz": "4", "color": "2980B9"},
            right={"sz": "4", "color": "2980B9"},
        )
        p2 = cell2.paragraphs[0]
        p2.paragraph_format.space_before = Pt(16)
        p2.paragraph_format.space_after = Pt(16)

        # Fleche entre cases
        if i < nb_cases - 1:
            fleche_idx = col_idx + 1
            p_f = table.rows[1].cells[fleche_idx].paragraphs[0]
            p_f.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_formatted_run(p_f, "\u2192", bold=True, size=16, color=BLEU_MOYEN)

    doc.add_paragraph()


def _build_lignes_reponse(doc, question, nb_lignes=3):
    """Construit des lignes pointillees pour la reponse."""
    espace = question.get("espace_reponse", "")
    if espace == "1_ligne":
        nb_lignes = 1
    elif espace == "2_lignes":
        nb_lignes = 2
    elif espace == "tableau":
        return  # Le tableau est deja construit

    _add_dotted_lines(doc, nb_lignes)


def _add_dotted_lines(doc, nb_lignes):
    """Ajoute des lignes pointillees."""
    for _ in range(nb_lignes):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run("." * 120)
        run.font.color.rgb = GRIS_CLAIR
        run.font.size = Pt(9)


def _build_corrige_answer(doc, question):
    """Construit la reponse corrigee d'une question."""
    corrige = question.get("corrige", {})
    if not corrige:
        return

    # Encadre reponse
    table = doc.add_table(rows=1, cols=1)
    cell = table.rows[0].cells[0]
    set_cell_shading(cell, FOND_HEX["vert_pale"])
    set_cell_border(
        cell,
        top={"sz": "8", "color": "27AE60"},
        bottom={"sz": "8", "color": "27AE60"},
        left={"sz": "8", "color": "27AE60"},
        right={"sz": "8", "color": "27AE60"},
    )

    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(4)
    add_formatted_run(p, "Reponse attendue : ", bold=True, size=10, color=VERT)
    add_formatted_run(p, corrige.get("reponse_attendue", ""), size=10)

    if corrige.get("explication"):
        p2 = cell.add_paragraph()
        p2.paragraph_format.space_before = Pt(4)
        p2.paragraph_format.space_after = Pt(4)
        add_formatted_run(p2, "Explication : ", bold=True, size=9, color=GRIS)
        add_formatted_run(p2, corrige["explication"], italic=True, size=9, color=GRIS)

    # Indicateurs
    indicateurs = question.get("indicateurs", {})
    niveaux = question.get("niveaux_possibles", [])
    if indicateurs and niveaux:
        p3 = cell.add_paragraph()
        p3.paragraph_format.space_before = Pt(4)
        p3.paragraph_format.space_after = Pt(2)
        add_formatted_run(p3, "Indicateurs : ", bold=True, size=9, color=BLEU_FONCE)
        for niv in niveaux:
            desc = indicateurs.get(niv, "")
            if desc:
                p4 = cell.add_paragraph()
                p4.paragraph_format.left_indent = Cm(0.5)
                p4.paragraph_format.space_before = Pt(1)
                p4.paragraph_format.space_after = Pt(1)
                add_formatted_run(p4, f"{niv} : ", bold=True, size=9,
                                  color=ROUGE if niv == "NT" else
                                  RGBColor(0xF3, 0x9C, 0x12) if niv == "I" else
                                  BLEU_MOYEN if niv == "A" else VERT)
                add_formatted_run(p4, desc, size=9)

    doc.add_paragraph()


def build_partie(doc, partie, is_corrige=False):
    """Construit une partie complete du devoir (1 approche)."""
    # Titre de partie
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(4)

    # Barre coloree + titre
    add_formatted_run(p, f"PARTIE {partie['numero']} : ", bold=True, size=14,
                      color=BLEU_FONCE)
    add_formatted_run(p, partie["titre"], bold=True, size=14, color=BLEU_MOYEN)

    # Sous-titre
    if partie.get("sous_titre"):
        add_paragraph_formatted(
            doc, partie["sous_titre"],
            italic=True, size=11, color=GRIS, space_after=8
        )

    # Ligne separatrice
    p_sep = doc.add_paragraph()
    run = p_sep.add_run("_" * 80)
    run.font.color.rgb = BLEU_CLAIR
    run.font.size = Pt(6)

    # Situation
    if partie.get("situation"):
        build_situation(doc, partie["situation"])

    # Documents
    for document in partie.get("documents", []):
        build_document_ref(doc, document)

    # Questions
    for question in partie.get("questions", []):
        build_question(doc, question, is_corrige=is_corrige)


def build_grille_notation(doc, grille, meta):
    """Construit la grille de notation par competences."""
    doc.add_page_break()

    add_paragraph_formatted(
        doc, "GRILLE D'EVALUATION PAR COMPETENCES",
        bold=True, size=16, color=BLEU_FONCE,
        alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=8
    )

    # Rappel titre devoir
    add_paragraph_formatted(
        doc, meta["titre"],
        bold=True, size=12, color=BLEU_MOYEN,
        alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=4
    )

    # Tableau eleve (mini)
    t_eleve = doc.add_table(rows=1, cols=3)
    t_eleve.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, label in enumerate(["Nom :", "Prenom :", "Classe :"]):
        cell = t_eleve.rows[0].cells[i]
        p = cell.paragraphs[0]
        add_formatted_run(p, label, bold=True, size=10, color=BLEU_FONCE)
        add_formatted_run(p, " " * 20, size=10)

    doc.add_paragraph()

    # Legende
    add_paragraph_formatted(
        doc, "Legende des niveaux de maitrise :",
        bold=True, size=10, color=BLEU_FONCE, space_before=8, space_after=4
    )

    legende = grille.get("niveaux_legende", {})
    for code, desc in legende.items():
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(1)
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(1)
        couleur = (ROUGE if code == "NT" else
                   RGBColor(0xF3, 0x9C, 0x12) if code == "I" else
                   BLEU_MOYEN if code == "A" else VERT)
        add_formatted_run(p, f"{code}", bold=True, size=10, color=couleur)
        add_formatted_run(p, f" = {desc}", size=10)

    doc.add_paragraph()

    # Tableau grille
    competences = grille.get("competences", [])
    cols = ["Competence", "Intitule", "Questions", "Points", "NT", "I", "A", "M", "Note"]
    table = doc.add_table(rows=1 + len(competences) + 1, cols=len(cols))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # En-tetes
    for i, h in enumerate(cols):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, FOND_HEX["bleu_fonce"])
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_formatted_run(p, h, bold=True, size=9, color=BLANC)

    # Lignes competences
    for idx, comp in enumerate(competences):
        row = table.rows[idx + 1]
        code = comp.get("code", "")
        comp_color = COULEURS_COMPETENCES.get(code, NOIR)

        # Code competence
        p = row.cells[0].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_formatted_run(p, code, bold=True, size=10, color=comp_color)

        # Intitule
        p = row.cells[1].paragraphs[0]
        add_formatted_run(p, comp.get("intitule", ""), size=8)

        # Questions associees
        questions = comp.get("questions_associees", [])
        p = row.cells[2].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_formatted_run(p, ", ".join(questions), size=8)

        # Points max
        p = row.cells[3].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_formatted_run(p, f'/ {comp.get("points_max", "")}', bold=True, size=10)

        # Cases NT/I/A/M (vides pour cocher)
        for col_idx in range(4, 8):
            p = row.cells[col_idx].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            add_formatted_run(p, "\u2610", size=14)

        # Note (vide)
        p = row.cells[8].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_formatted_run(p, "/ " + str(comp.get("points_max", "")), size=9,
                          color=GRIS)

        # Alternance couleur
        if idx % 2 == 0:
            for c in row.cells:
                set_cell_shading(c, FOND_HEX["bleu_pale"])

    # Ligne total
    total_row = table.rows[-1]
    set_cell_shading(total_row.cells[0], FOND_HEX["bleu_fonce"])

    # Fusionner les premieres cellules pour le total
    for i in range(4):
        set_cell_shading(total_row.cells[i], FOND_HEX["bleu_fonce"])
    p = total_row.cells[0].paragraphs[0]
    add_formatted_run(p, "", size=1)
    p = total_row.cells[3].paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_formatted_run(p, "TOTAL", bold=True, size=11, color=BLANC)

    for i in range(4, 8):
        set_cell_shading(total_row.cells[i], FOND_HEX["bleu_fonce"])
        p = total_row.cells[i].paragraphs[0]
        add_formatted_run(p, "", size=1)

    # Case note finale
    p = total_row.cells[8].paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)
    total = grille.get("total", 20)
    add_formatted_run(p, f"/ {total}", bold=True, size=14, color=BLEU_FONCE)

    doc.add_paragraph()

    # Note methodologique
    methode = grille.get("methode", "")
    if methode:
        add_paragraph_formatted(
            doc, "Methode de notation :",
            bold=True, size=10, color=BLEU_FONCE, space_before=8
        )
        add_paragraph_formatted(
            doc, methode,
            italic=True, size=9, color=GRIS, space_after=8
        )


def build_corrige(doc, data):
    """Construit le document corrige complet."""
    meta = data["meta"]

    doc.add_page_break()

    # Bandeau CORRIGE
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.rows[0].cells[0]
    set_cell_shading(cell, "FDEDEC")
    set_cell_border(
        cell,
        top={"sz": "12", "color": "E74C3C"},
        bottom={"sz": "12", "color": "E74C3C"},
        left={"sz": "12", "color": "E74C3C"},
        right={"sz": "12", "color": "E74C3C"},
    )
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)
    add_formatted_run(p, "CORRIGE - DOCUMENT ENSEIGNANT", bold=True, size=16,
                      color=ROUGE)

    add_paragraph_formatted(
        doc, meta["titre"],
        bold=True, size=13, color=BLEU_MOYEN,
        alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=8, space_after=12
    )

    # Parties avec reponses
    for partie in data.get("parties", []):
        build_partie(doc, partie, is_corrige=True)


def generate_document(data, output_path=None):
    """Genere le document Word complet a partir du JSON."""
    meta = data["meta"]
    doc = Document()

    # Configuration des marges
    for section in doc.sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)

    # EN-TETE
    build_header(doc, meta)

    # PARTIES (3 approches)
    for partie in data.get("parties", []):
        build_partie(doc, partie, is_corrige=False)

    # GRILLE DE NOTATION
    if data.get("grille_notation"):
        build_grille_notation(doc, data["grille_notation"], meta)

    # CORRIGE (optionnel)
    if meta.get("avec_corrige", False):
        build_corrige(doc, data)

    # Sauvegarde
    if not output_path:
        module = meta.get("module", "PSE")
        niveau = meta.get("type_diplome", "").replace(" ", "_")
        diff = meta.get("difficulte", "moyen")
        output_path = f"Devoir_PSE_{module}_{niveau}_{diff.capitalize()}.docx"

    doc.save(output_path)
    return output_path


# =============================================================================
# POINT D'ENTREE
# =============================================================================

def main():
    if len(sys.argv) < 2:
        print("Usage: python generer_devoir.py <fichier.json> [--output <fichier.docx>]")
        print()
        print("Exemple:")
        print("  python generer_devoir.py devoir_A3_cap.json")
        print("  python generer_devoir.py devoir_A3_cap.json --output mon_devoir.docx")
        sys.exit(1)

    json_path = sys.argv[1]

    output_path = None
    if "--output" in sys.argv:
        idx = sys.argv.index("--output")
        if idx + 1 < len(sys.argv):
            output_path = sys.argv[idx + 1]

    if not os.path.exists(json_path):
        print(f"Erreur: Le fichier '{json_path}' n'existe pas.")
        sys.exit(1)

    with open(json_path, "r", encoding="utf-8") as f:
        data = json.load(f)

    result = generate_document(data, output_path)
    print(f"Document genere : {result}")


if __name__ == "__main__":
    main()
